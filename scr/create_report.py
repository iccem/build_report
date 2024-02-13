from docx import Document
from docx.shared import Mm
import imgs
import sqlite3


def _get_report_links() -> list:
    con = sqlite3.connect('scr/doc/online_adv.db')
    cursor = con.cursor()
    temp = []
    links_report_list = []

    cursor.execute('''select * from STG_REPORT_LINKS''')

    for row in cursor.fetchall():
        temp.append(row)

    for row in temp:
        date = str(row[1]).replace(" 00:00:00", "")
        path = row[3]
        links_report_list.append(date + '\n' + path)
    return links_report_list


def print_imgs_by_date(current_date, pics, document, path_imgs):
    paragraph = document.add_paragraph('')
    for i in pics:
        p = path_imgs + i
        # print(p)
        if current_date in p:
            try:
                run = paragraph.add_run()
                inline_shape = run.add_picture(p)
                inline_shape.width = Mm(60)
                inline_shape.height = Mm(86)
            except ValueError:
                print("Oops!  That was no valid number.  Try again...")

def get_report(company, month, path_imgs):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(8)
        section.left_margin = Mm(12)
        section.right_margin = Mm(12)

    document.add_heading('ОТЧЕТ О РАЗМЕЩЕНИИ. Проект: «ОНЛАЙН»', 2)
    # document.add_heading('Проект: «ОНЛАЙН»', 2)
    document.add_heading(f'Период: {month}', 3)

    # list of links
    report_links = _get_report_links()
    current_date_print = ''

    pics = imgs.get_pics_list(path_imgs)
    
    paragraph = document.add_paragraph('')

    for link in report_links:
        current_date, current_link = link.split('\n')
        if current_date_print != current_date:
            p = document.add_paragraph('\n')
            p.add_run(current_date).bold = True
            print_imgs_by_date(current_date, pics, document, path_imgs)
            p = document.add_paragraph(current_link)
            current_date_print = current_date
        elif current_date_print == current_date:
            p = document.add_paragraph(current_link)
            

    # document.add_page_break()

    document.save(f'{company} {month}.docx')